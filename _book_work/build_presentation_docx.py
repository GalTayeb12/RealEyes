# -*- coding: utf-8 -*-
"""Build comprehensive Hebrew Word presentation guide with images."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/home/sceuser/RealEyes/gdrive/deepfake_image_project/models/madrich_hearatza_RealEyes.docx'
VIZ = '/home/sceuser/RealEyes/gdrive/deepfake_image_project/models/RealEyes_v2/visualizations'
LORA = '/home/sceuser/RealEyes/gdrive/deepfake_image_project/models/lora_vs_no_lora'

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────
def rtl(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn('w:bidi')) is None:
        b = OxmlElement('w:bidi'); b.set(qn('w:val'), '1'); pPr.append(b)
    for r in p.runs:
        rPr = r._r.get_or_add_rPr()
        if rPr.find(qn('w:rtl')) is None:
            e = OxmlElement('w:rtl'); e.set(qn('w:val'), '1'); rPr.append(e)

def h1(t):
    p = doc.add_heading(t, level=1); rtl(p)

def h2(t):
    p = doc.add_heading(t, level=2); rtl(p)

def h3(t):
    p = doc.add_heading(t, level=3); rtl(p)

def body(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(11)
    rtl(p); return p

def bullet(t):
    p = doc.add_paragraph(t, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
    rtl(p)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    try: t.style = 'Table Grid'
    except Exception: pass
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

def add_fig(path, caption, explanation, width=6.0):
    if not os.path.exists(path):
        body(f'[תמונה לא נמצאה: {path}]'); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cp = doc.add_paragraph()
    cr = cp.add_run(caption); cr.bold = True; cr.font.size = Pt(10)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(cp)
    body(explanation)
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════
# TITLE
# ═════════════════════════════════════════════════════════════════════
t = doc.add_heading('מדריך הכנה מקיף להצגת פרויקט RealEyes', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
body('מסמך זה מסביר את כל המחברות, המודלים, הניסויים, הויזואליזציות והמסקנות — ברמה שמאפשרת להציג את הפרויקט גם למי שלא ראה את הקוד.')
body('גרסה: יוני 2026 | מחברת ייצור: the_modles_fixed_new.ipynb | ניסויים: exp_* (LoRA / no-LoRA)')

# ═════════════════════════════════════════════════════════════════════
h1('חלק א\' — רקע ויעדים')
# ═════════════════════════════════════════════════════════════════════
body('RealEyes היא מערכת לזיהוי תמונות Deepfake (מזויפות מול אמיתיות). הגישה שלנו היברידית:')
bullet('ניתוח רעש/תדר (SRM) — חושף "טביעות אצבע" של מודלים גנרטיביים')
bullet('ניתוח ויזואלי (EfficientNet) — לומד מאפיינים סמנטיים (פנים, טקסטורה)')
bullet('שילוב Two-Stream + Ensemble — מיזוג שני העולמות')
bullet('מחקר Cross-Database + LoRA — בדיקת הכללה וייעול אימון')

h2('יעדים שהגדרנו')
add_table(['יעד','ערך מבוקש','מה השגנו'],
[['Accuracy (Ensemble)','> 85%','61.6–62.6%'],
 ['ROC-AUC (Ensemble)','> 0.90','0.745–0.752'],
 ['יציבות Cross-DB','פער < 10%','פער 20–31%'],
 ['זמן תגובה (אתר)','< 3 שניות','EfficientNet בודד — מהיר']])

h2('למה לא הגענו ל-85% Accuracy?')
bullet('סף החלטה 0.5 לא אופטימלי — CNN-SRM כמעט תמיד חוזה REAL (Recall FAKE=3%), אך AUC=0.697 מראה יכולת הפרדה.')
bullet('חוסר איזון מחלקות — 10,504 FAKE מול 8,337 REAL; Class Weights מכוילים לPrecision גבוה.')
bullet('מגוון מאגרים — 4 מאגרים עם טכנולוגיות יצירה שונות (GAN, Diffusion, Face-swap).')
bullet('Cross-DB — Within-DB AUC 0.73–0.97, Cross-DB 0.49–0.56. האתגר האמיתי הוא הכללה, לא Overfitting.')
bullet('LoRA פגע ב-CNN-SRM על OpenForensics (ירידה 0.35 ב-AUC).')
body('מה לומר בהצגה: "AUC 0.75 מראה שהמערכת יודעת להפריד. Accuracy נמוך בגלל סף + הכללה — ממצא מחקרי, לא כישלון."', bold=True)

# ═════════════════════════════════════════════════════════════════════
h1('חלק ב\' — מחברת the_modles_fixed_new.ipynb: הסבר תא-אחר-תא')
# ═════════════════════════════════════════════════════════════════════
body('המחברת בנויה כ"צינור" (Pipeline): כל תא תלוי בתוצרי התאים שלפניו. להלן הסבר לכל תא, כאילו את רואה את הקוד בפעם הראשונה.')

CELLS = [
(0, 'Markdown — כותרת', 'כותרת המחברת: "Testing classifiers and models". אין קוד — רק תיאור.'),
(1, 'Setup — ייבוא ספריות והגדרות GPU', '''מה קורה: מייבאים os, numpy, tensorflow וכו'.
למה כיבינו XLA/JIT: מניעת שגיאות BatchNormalization ("JIT compilation failed").
Memory Growth=True: GPU לא תופס את כל הזיכרון מראש — חשוב כשרצים כמה מודלים.
Seed=42: reproducibility — אותה תוצאה בכל הרצה.'''),
(2, 'Paths — נתיבים למאגרים ומודלים', '''מגדיר איפה הכל נמצא:
GDRIVE_PATH = ~/RealEyes/gdrive
MODELS_DIR = .../models/RealEyes_v2 (שמירת מודלים)
DATA_PATH, DATA_SPLIT_PATH — CustomWar
datasets/ — OpenForensics, CiFake, CelebDF
SKIP_IF_SAVED=True — אם מודל כבר שמור, לא מאמנים מחדש (חוסך שעות).'''),
(3, 'Sync Drive → Server (rclone)', 'מסנכרן תמונות ומודלים מ-Google Drive לשרת. rclone מדלג על קבצים זהים — בטוח להריץ שוב.'),
(4, 'Markdown — Celeb-DF Setup', 'הסבר: CelebDF v2 הוא מאגר וידאו — צריך לחלץ פריימים (פעם אחת).'),
(5, 'Celeb-DF Frame Extraction — DO NOT RUN', 'מחלץ פריים לשנייה מסרטוני CelebDF. מסומן DO NOT RUN — כבר בוצע. לא להריץ שוב.'),
(6, 'Auto Split CustomWar — DO NOT RUN', 'מחלק CustomWar ל-train/val/test (80/10/10). DO NOT RUN — כבר בוצע.'),
(7, 'Helper: load_dataset_images', '''פונקציה שטוענת נתיבי תמונות + תוויות (0=REAL, 1=FAKE).
בודקת סיומות (.jpg/.png), מדלגת על קבצים פגומים.
REAL/Fake נקבע לפי שם תיקייה או מבנה המאגר.'''),
(8, 'Build FULL Train / Val / Test Sets', '''טוען את כל 4 המאגרים ל-3 סטים:
Train: ~318K תמונות | Val: ~68K | Test: ~18,841
כל מאגר (OpenForensics, CiFake, CustomWar, CelebDF) נטען בנפרד ואז מאוחד.'''),
(9, 'Dataset Statistics', 'מדפיס סיכום: כמה REAL ו-FAKE בכל סט. עוזר לזהות חוסר איזון.'),
(10, 'Shared Preprocessing — פרמטרים גלובליים', '''IMG_SIZE_RGB=(224,224) — לEfficientNet
IMG_SIZE_SRM=(256,256) — לערוץ הרעש
BATCH_SIZE_RGB=8, SRM=32, HYBRID=4 — HYBRID קטן כי Two-Stream = 2 מודלים בזיכרון.
AUTOTUNE — TensorFlow בוחר מקביליות אופטימלית.'''),
(11, 'Markdown — SRM Model', 'מתחילים את שלב SRM (Spatial Rich Model) — ניתוח רעש.'),
(12, 'SRM Filters — 8 פילטרים 5×5', '''מגדירים 8 פילטרים קבועים (לא נלמדים!) שמחלצים:
Laplacian, Edge, Diagonal, Vertical, High-Frequency...
+ 3 פילטרים 3×3 שמורחבים ל-5×5.
הרעיון: GAN/Diffusion משאירים "חתימות" בדומיין הרעש.'''),
(13, 'apply_srm_filters_tf()', 'פונקציית TensorFlow שמפעילה conv2d עם 8 הפילטרים → 8 ערוצי feature maps.'),
(14, 'SRM Dataset Pipeline', '''יוצר tf.data.Dataset:
decode → resize 256×256 → apply SRM → augment (flip, brightness) → batch → prefetch.
Augmentation רק ב-train — לא ב-val/test.'''),
(15, 'Prepare FULL SRM Datasets', 'יוצר train_srm_dataset, val_srm_dataset, test_srm_dataset מהנתיבים שנטענו בתא 8.'),
(16, 'SRM Sanity Check', 'בודק batch אחד: shape, min/max — לוודא שהpipeline תקין לפני אימון.'),
(17, 'Clear GPU Memory', 'gc.collect() + clear_session() — מנקה זיכרון GPU לפני בניית מודל חדש. מציג nvidia-smi.'),
(18, 'Improved Autoencoder / Encoder', '''Autoencoder משופר:
SE blocks — לומד אילו ערוצי רעש חשובים.
Residual connections — gradient flow טוב יותר.
L2 reg (1e-4) — anti-overfitting.
Input: SRM 24ch → Encoder: (16,16,256) → Decoder → reconstruction.
Loss: MSE — לומד לדחוס/לשחזר מפת רעש.'''),
(19, 'Autoencoder Sanity Check', 'מריץ batch אחד דרך AE — בודק shapes. try/except — שגיאת GPU לא תעצור את המחברת.'),
(20, 'Prepare AE Datasets', 'ae_train_ds: (x,x) — input=output (reconstruction). ae_val_ds: validation.'),
(21, 'Autoencoder Training', '''מאמן AE אם SKIP_IF_SAVED=False.
Callbacks: EarlyStopping(val_loss, patience=5), ModelCheckpoint.
אם SKIP_IF_SAVED=True — טוען AE שמור מ-Drive.'''),
(22, 'Save / Load Encoder', 'שומר encoder + autoencoder ל-Drive. Encoder ישמש את CNN-SRM.'),
(23, 'Clear Memory Before Classifier', 'ניקוי זיכרון לפני בניית CNN-SRM.'),
(24, 'CNN-SRM Classifier — ארכיטקטורה', '''Encoder (frozen Phase 1) → ConvBlock×2 → GAP → Dense(1,sigmoid).
ConvBlock = Conv2D + BN + ReLU + MaxPool.
Phase 1: encoder.trainable=False — רק Head לומד.
Phase 2: 70% תחתון encoder קפוא, 30% עליון unfrozen.'''),
(25, 'CNN-SRM Sanity Check', 'batch אחד → prediction shape — לוודא שהמודל עובד.'),
(26, 'Class Weights', '''compute_class_weight("balanced") → {REAL: 1.202, FAKE: 0.856}
REAL מקבל משקל גבוה יותר — כי יש פחות דוגמאות REAL.
מועבר ל-model.fit(class_weight=...) — מונע התעלמות מ-REAL.'''),
(27, 'CNN-SRM Phase 1 Training', '''אם SKIP_IF_SAVED: טוען final_cnn_srm.keras — מדלג.
אחרת: fit עם LR=1e-4, EarlyStopping(val_auc, patience=4), ReduceLROnPlateau.
Loss: BinaryCrossentropy(label_smoothing=0.1) — מונע overconfidence.'''),
(28, 'CNN-SRM Phase 2 Fine-tune', '''Unfreeze 30% עליון encoder. LR=1e-5 (נמוך!), clipnorm=1.0.
Phase 2 = fine-tune עדין — לא להרוס features של SRM.'''),
(29, 'CNN-SRM Test Evaluation', 'model.evaluate(test_srm_dataset) — Accuracy + AUC על Test.'),
(30, 'CNN-SRM Detailed Report', '''Classification Report + ROC-AUC.
תוצאה: Accuracy 45.4%, AUC 0.697, Recall FAKE=3.2%.
המודל נוטה ל-REAL — AUC סביר אבל Accuracy נמוך.'''),
(31, 'Save CNN-SRM Model', 'שומר final_cnn_srm.keras ל-Drive.'),
(32, 'CNN-SRM Visualizations', 'Confusion Matrix + ROC + Score Histogram → cnn_srm_results.png'),
(33, 'Markdown — EfficientNetB0', 'מתחילים מודל RGB ויזואלי.'),
(34, 'EfficientNet Imports + Class Weights', 'ייבוא EfficientNetB0, preprocess_input, class weights (אותם כמו SRM).'),
(35, 'build_efficientnet_b0()', '''EfficientNetB0(ImageNet) → GAP → Dropout(0.3) → Dense(256) → Dropout(0.25) → Dense(64) → Dropout(0.15) → sigmoid.
train_base=False בStage 1 — backbone קפוא.'''),
(36, 'Class Weights (duplicate check)', 'מוודא class weights מוגדרים.'),
(37, 'EfficientNet Stage 1', '''Backbone frozen, LR=5e-4, fit על train_rgb_ds.
Callbacks: EarlyStopping(val_auc), ModelCheckpoint(stage1).'''),
(38, 'EfficientNet Stage 2 Fine-tune', '''Unfreeze top 30% EfficientNet. LR=1e-5.
טוען stage1 → fine-tune → שומר efficientnetb0_finetuned_best.keras.'''),
(39, 'EfficientNet Test Evaluation', 'evaluate על test_rgb_ds.'),
(40, 'EfficientNet Detailed Report', 'Accuracy 62.6%, AUC 0.715 — המאוזן ביותר.'),
(41, 'EfficientNet Visualizations', 'CM + ROC + Histogram → efficientnet_results.png'),
(42, 'Markdown — Two-Stream', 'שילוב SRM + RGB בשני backbones.'),
(43, 'Two-Stream Dataset Pipeline', '''כל sample = (srm_batch, rgb_batch), label.
SRM: 256×256×24 | RGB: 224×224×3 (EfficientNet preprocess).
Batch=4 — זיכרון כפול.'''),
(44, 'Two-Stream Architecture', '''טוען encoder + EfficientNet (frozen) → Concatenate features → Dense head → sigmoid.
שני streams על אותה תמונה — noise + appearance.'''),
(45, 'Two-Stream Phase 1', 'Head only, backbones frozen.'),
(46, 'Two-Stream Phase 2', 'Partial unfreeze 30% עליון.'),
(47, 'Two-Stream Evaluation', 'Accuracy 66.1%, AUC 0.722 — הטוב ביותר בודד.'),
(48, 'Two-Stream Visualizations', '→ two_stream_results.png'),
(49, 'Markdown — Ensemble', 'שילוב 3 מודלים: SRM + EfficientNet + Two-Stream.'),
(50, 'Collect Predictions (Val + Test)', '''מריץ predict על כל 3 המודלים → val_p_srm/eff/ts, test_p_srm/eff/ts.
אלה ההסתברויות (0–1) לכל תמונה — בסיס ל-Ensemble.'''),
(51, 'Uniform Ensemble (1/3 each)', 'ממוצע פשוט → Accuracy 62.6%, AUC 0.752.'),
(52, 'Optimal Ensemble (Nelder-Mead)', '''מחפש משקלים שממקסמים val AUC:
SRM: 0.445 | EfficientNet: 0.483 | TwoStream: 0.072
Test: Accuracy 61.6%, AUC 0.745. Two-Stream כמעט לא תורם — EfficientNet מכסה.'''),
(53, 'Ensemble + All Models Comparison Viz', 'ensemble_results.png, all_models_comparison.png, model_ranking.png...'),
(54, 'Markdown — Grad-CAM', 'הסבר החלטות — Heatmap על התמונה.'),
(55, 'Grad-CAM Utilities', 'פונקציות core: gradient של output לפי conv layer → heatmap.'),
(56, 'Grad-CAM Prediction Functions', 'gradcam_efficientnet(path), gradcam_cnn_srm(path) — לשימוש באתר.'),
(57, 'Grad-CAM Batch Visualization', 'דוגמאות REAL/FAKE עם heatmaps — שמירה ל-Drive.'),
(58, 'Grad-CAM Demo', 'הרצה על תמונה בודדת — demo.'),
(59, 'Sync Visualizations to Drive', 'rclone sync → gdrive:.../visualizations/'),
]

for idx, title, expl in CELLS:
    h3(f'תא {idx}: {title}')
    body(expl)

# ═════════════════════════════════════════════════════════════════════
h1('חלק ג\' — ניסויים Cross-Database (exp_*.ipynb)')
# ═════════════════════════════════════════════════════════════════════
body('שלוש מחברות נפרדות בודקות הכללה: מה קורה כש**מאמנים** על מאגר A ו**בודקים** על מאגר B?')

h2('מתודולוגיה')
bullet('4 מאגרים: OpenForensics, CiFake, CustomWar, CelebDF + Combined (כולם יחד)')
bullet('לכל מאגר אימון: 4+4=16 קומבינציות Cross-DB')
bullet('מדדים: ROC-AUC, Accuracy, F1-FAKE')
bullet('Within-DB = אלכסון (train=test same DB) | Cross-DB = off-diagonal')

h2('מחברות')
add_table(['מחברת','LoRA','תיאור'],
[['exp_cnn_srm.ipynb','כן','CNN-SRM + LoRADense'],
 ['exp_cnn_srm_no_lora.ipynb','לא','Full Fine-Tuning'],
 ['exp_efficientnet.ipynb','כן','EfficientNetB0 + LoRA'],
 ['exp_efficientnet_no_lora.ipynb','לא','Full FT'],
 ['exp_vit.ipynb','כן','ViT-Tiny + LoRA בכל MLP'],
 ['exp_vit_no_lora.ipynb','לא','Full FT']])

h3('פרמטרים משותפים בניסויים')
add_table(['פרמטר','CNN-SRM','EfficientNet','ViT'],
[['Batch','32','32 (16 Combined)','32'],
 ['Phase 1 LR','5e-4','5e-4','AdamW 5e-5'],
 ['Phase 2 LR','1e-5','1e-5','AdamW 1e-5'],
 ['LoRA rank/alpha','16/32','16/32','16/32'],
 ['Patience','4-5','5','8-10'],
 ['IMG size','256 SRM','224 RGB','224 RGB']])

h2('תוצאות Cross-DB — סיכום AUC (LoRA)')
add_table(['מודל','Within avg','Cross avg','פער'],
[['CNN-SRM','0.731','0.493','23.7%'],
 ['EfficientNetB0','0.829','0.522','30.7%'],
 ['ViT-Light','0.731','0.525','20.7%']])

h2('LoRA vs Full Fine-Tuning')
add_table(['מודל','Pre-LoRA Within','LoRA Within','Δ'],
[['CNN-SRM','0.852','0.731','-0.121'],
 ['EfficientNet','0.895','0.829','-0.066'],
 ['ViT','0.695','0.743','+0.047']])
body('מסקנה: LoRA יעיל ל-ViT ולEdge devices. CNN-SRM נפגע — במיוחד OpenForensics (0.871→0.522).')

h2('אופטימיזציות זיכרון (למה הקטנו)')
bullet('shuffle buffer: 2000 (לא כל הדאטה)')
bullet('prefetch(4) במקום AUTOTUNE')
bullet('ram_budget = 3GB')
bullet('gc.collect() + malloc_trim(0) בין מאגרים')
bullet('Batch=16 ל-Combined dataset')
bullet('try-except סביב model.fit — שרת לא קורס')

# ═════════════════════════════════════════════════════════════════════
h1('חלק ד\' — ויזואליזציות: מה רואים ומה מסיקים')
# ═════════════════════════════════════════════════════════════════════

VIZ_EXPLAIN = [
(VIZ+'/cnn_srm_results.png', 'איור 1: תוצאות CNN-SRM',
 '''שלושה גרפים:
1) Confusion Matrix — רוב התחזיות ב-TN (REAL נכון) ו-FN (FAKE שפוספס) → המודל אומר REAL כמעט תמיד.
2) ROC Curve — AUC=0.697 — יש הפרדה, לא אקראי.
3) Score Histogram — רוב הציונים נמוכים (קרוב ל-0=REAL).
מסקנה: המודל שמרני מדי; ב-Ensemble תורם ל-Precision.'''),
(VIZ+'/efficientnet_results.png', 'איור 2: תוצאות EfficientNetB0',
 '''CM: TP=4968 FAKE נכון, TN=6827 REAL נכון — מאוזן.
ROC: AUC=0.715.
Histogram: שני peaks — REAL (~0.2) ו-FAKE (~0.7).
מסקנה: המודל הטוב ביותר לשימוש יחיד — **זה המודל באתר**.'''),
(VIZ+'/two_stream_results.png', 'איור 3: תוצאות Two-Stream',
 '''Accuracy 66.1%, AUC 0.722 — הטוב ביותר בודד.
Recall FAKE=54% — מזהה יותר FAKE מ-EfficientNet.
מסקנה: שילוב SRM+RGB עוזר, אך ב-Ensemble משקל נמוך (7%).'''),
(VIZ+'/ensemble_results.png', 'איור 4: Ensemble אופטימלי',
 '''Accuracy 61.6%, AUC 0.745.
Precision FAKE=79.4% — כשאומר FAKE, צודק ב-79%.
Recall FAKE=42% — מפספס חלק מהזיופים.
מסקנה: מכוון לPrecision — מתאים ליישום שרוצה מעט False Positives.'''),
(VIZ+'/all_models_comparison.png', 'איור 5: השוואת כל המודלים',
 '''שמאל: ROC curves — Ensemble (uniform) הכי גבוה (0.752).
ימין: Bar chart — Accuracy/ROC-AUC/Precision/Recall/F1 לכל מודל.
Two-Stream מוביל ב-Accuracy; Ensemble ב-AUC.
מסקנה: אין "מנצח אחד" — Ensemble משלב חוזקות.'''),
(VIZ+'/all_models_score_table.png', 'איור 6: טבלת ציונים',
 'טבלה צבעונית — ירוק=הטוב ביותר, אדום=החלש. מראה במבט אחד מי מוביל בכל מדד.'),
(VIZ+'/model_ranking.png', 'איור 7: דירוג לפי ROC-AUC',
 'Ensemble (uniform) ראשון, Two-Stream שני, EfficientNet שלישי, CNN-SRM אחרון.'),
(VIZ+'/real_vs_fake_detection.png', 'איור 8: שיעור זיהוי REAL vs FAKE',
 '''כל מודל — כמה REAL ו-FAKE זוהו נכון.
CNN-SRM: כמעט כל REAL, מעט FAKE.
EfficientNet/Two-Stream: יותר מאוזן.
מסקנה: CNN-SRM "זהיר" — לא מתריע FAKE בקלות.'''),
(LORA+'/within_db_auc_comparison.png', 'איור 9: השוואת Within-DB AUC — LoRA vs Full FT',
 'עמודות: Pre-LoRA (אימון מלא) לעומת LoRA. ViT עולה; CNN-SRM/EfficientNet יורדים מעט.'),
(LORA+'/within_db_delta.png', 'איור 10: Δ AUC — השפעת LoRA',
 'ירוק=שיפור, אדום=ירידה. ViT-CiFake/Combined ירוק; CNN-OF אדום חזק.'),
(LORA+'/cnn_srm_cross_db_annotated.png', 'איור 11: מטריצת CNN-SRM Cross-DB',
 'Heatmap AUC. אלכסון=Within (גבוה). off-diagonal=Cross (נמוך). OF→* קריסה.'),
(LORA+'/efficientnet_cross_db_annotated.png', 'איור 12: מטריצת EfficientNet Cross-DB',
 'Within גבוה (0.8-1.0). Cross נמוך (0.4-0.6). Combined→CW=0.87.'),
(LORA+'/vit_cross_db_annotated.png', 'איור 13: מטריצת ViT Cross-DB',
 'הכללה הטובה ביותר. Combined→CiFake=0.92. פער Within-Cross הקטן ביותר.'),
(LORA+'/lora_cross_db_roc_auc.png', 'איור 14: Heatmaps AUC — כל 3 המודלים',
 '3 heatmaps זה לצד זה — השוואה ויזואלית של Cross-DB לכל ארכיטקטורה.'),
(LORA+'/summary_table.png', 'איור 15: טבלת סיכום LoRA vs no-LoRA',
 'טבלה מרכזת: Within, Cross, פער, מסקנה לכל מודל.'),
]

for path, cap, expl in VIZ_EXPLAIN:
    h3(cap)
    add_fig(path, cap, expl)

# ═════════════════════════════════════════════════════════════════════
h1('חלק ה\' — מסקנות סופיות ו-Q&A')
# ═════════════════════════════════════════════════════════════════════
CONCL = [
'בנינו מערכת היברידית: SRM + EfficientNet + Two-Stream + Ensemble.',
'AUC 0.75 (Ensemble) — יכולת הפרדה טובה. Accuracy 62-66% — יש מקום לשיפור (סף + הכללה).',
'Cross-DB הוא האתגר: פער 20-31% בין Within ל-Cross.',
'ViT+LoRA — הכללה הטובה ביותר. EfficientNet — Within הטוב ביותר.',
'Combined training משפר Cross-DB.',
'LoRA — יעיל לEdge; ViT נהנה, CNN-SRM נפגע על OF.',
'Xception ננטש — OOM. EfficientNetB0 = בחירה נכונה.',
'Grad-CAM — שקיפות. האתר משתמש ב-EfficientNet בלבד.',
'השערת מחקר (SRM יכליל הכי טוב) — נדחתה. ViT הכליל הכי טוב.',
'יעד 85% — לא הושג; ממצא מחקרי על הכללה חשוב יותר מ-Accuracy גבוה על מאגר בודד.',
]
body('10 מסקנות להצגה:', bold=True)
for i,c in enumerate(CONCL,1):
    bullet(f'{i}. {c}')

h2('שאלות צפויות')
QA = [
('למה Accuracy CNN-SRM רק 45%?','המודל נוטה ל-REAL. AUC=0.697 = יש הפרדה. ב-Ensemble תורם Precision.'),
('Within vs Cross?','Within=אותו מאגר (קל). Cross=מאגר אחר (קשה, ריאלי).'),
('מה LoRA?','כיוונון יעיל — 0.1% פרמטרים. ViT נהנה, CNN פחות.'),
('למה לא Xception?','OOM — 2×22M params. EfficientNet=5M.'),
('מה באתר?','EfficientNetB0 + Grad-CAM. AUC 0.715.'),
('למה לא Ensemble באתר?','3× inference time. EfficientNet מספיק.'),
]
for q,a in QA:
    body(f'ש: {q}', bold=True)
    body(f'ת: {a}')

doc.save(OUT)
print('SAVED:', OUT, 'size', os.path.getsize(OUT))
