import docx
from docx import Document

doc = Document('/home/sceuser/RealEyes/RealEyes/_book_work/project_book.docx')

print("=== PARAGRAPHS (idx | style | text) ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"{i}\t[{p.style.name}]\t{t[:120]}")

print("\n\n=== TABLES ===")
for ti, tbl in enumerate(doc.tables):
    print(f"\n--- TABLE {ti} : rows={len(tbl.rows)} cols={len(tbl.columns)} ---")
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip().replace('\n',' ')[:30] for c in row.cells]
        print(f"  r{ri}: {cells}")
